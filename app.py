# from flask import Flask, redirect, render_template, request, url_for
# import sqlite3
# import re
# from flask import flash # Add this import at the top
# from docx import Document
# from flask import send_file
# import io

# app = Flask(__name__)

# # 1. Change the main route to show the login page first
# @app.route('/')
# def home():
#     return render_template('login.html')

# @app.route('/login', methods=['GET', 'POST'])
# def login():
#     if request.method == 'POST':
#         email = request.form.get('email')
#         password = request.form.get('password')

#         conn = get_db_connection()
#         # Find the user by email
#         user = conn.execute('SELECT * FROM users WHERE email = ?', (email,)).fetchone()
#         conn.close()

#         if user and user['password'] == password:
#             # Login successful! 
#             # In a real app, you would use session['user_id'] = user['id'] here.
#             return redirect(url_for('index'))
#         else:
#             return "Invalid email or password. Please try again."
#         return redirect(url_for('dashboard'))
#     return render_template('login.html')

# # @app.route('/signup', methods=['GET', 'POST'])
# # def signup():
# #     return render_template('signup.html') # Create this similar to login.html
# @app.route('/signup', methods=['GET', 'POST'])
# def signup():
#     if request.method == 'POST':
#         fullname = request.form.get('fullname')
#         email = request.form.get('email')
#         password = request.form.get('password')
#         confirm_password = request.form.get('confirm_password')

#         # 1. Check if passwords match
#         if password != confirm_password:
#             return "Passwords do not match! Please go back and try again."

#         # 2. Basic Length Check (as per your MovieMagic UI)
#         if len(password) < 8:
#             return "Password must be at least 8 characters long."

#         conn = get_db_connection()
#         try:
#             conn.execute('INSERT INTO users (fullname, email, password) VALUES (?, ?, ?)',
#                          (fullname, email, password))
#             conn.commit()
#             conn.close()
#             # Redirect to login page upon success
#             return redirect(url_for('login')) 
            
#         except sqlite3.IntegrityError:
#             conn.close()
#             return "Email already exists! Try logging in instead."
            
#     return render_template('signup.html')
# @app.route('/dashboard')
# def dashboard():

# @app.route('/export-word')
# def export_word():
#     conn = get_db_connection()
#     records = conn.execute('SELECT s.name, e.event_name, e.date, e.venue FROM participation p JOIN students s ON p.student_id = s.id JOIN events e ON p.event_id = e.id').fetchall()
#     conn.close()

#     # Create Word Document
#     doc = Document()
#     doc.add_heading('BITM CSE - Hackathon Participation Report', 0)

#     table = doc.add_table(rows=1, cols=4)
#     table.style = 'Table Grid'
#     hdr_cells = table.rows[0].cells
#     hdr_cells[0].text = 'Student Name'
#     hdr_cells[1].text = 'Event Name'
#     hdr_cells[2].text = 'Date'
#     hdr_cells[3].text = 'Venue'

#     for rec in records:
#         row_cells = table.add_row().cells
#         row_cells[0].text = str(rec['name'])
#         row_cells[1].text = str(rec['event_name'])
#         row_cells[2].text = str(rec['date'])
#         row_cells[3].text = str(rec['venue'])

#     # Save to memory and send to browser
#     file_stream = io.BytesIO()
#     doc.save(file_stream)
#     file_stream.seek(0)
    
#     return send_file(file_stream, as_attachment=True, download_name='Hackathon_Report.docx')



# @app.route('/delete-record/<int:id>')
# def delete_record(id):
#     conn = get_db_connection()
#     conn.execute('DELETE FROM participation WHERE id = ?', (id,))
#     conn.commit()
#     conn.close()
    
#     flash("deleted") # Send a signal to the frontend
#     return redirect(url_for('index'))

# def get_db_connection():
#     conn = sqlite3.connect('hackathon.db')
#     conn.row_factory = sqlite3.Row # Allows accessing columns by name
#     return conn
# def init_users_db():
#     conn = sqlite3.connect('hackathon.db')
#     # This command creates the table only if it's missing
#     conn.execute('''
#         CREATE TABLE IF NOT EXISTS users (
#             id INTEGER PRIMARY KEY AUTOINCREMENT,
#             fullname TEXT NOT NULL,
#             email TEXT UNIQUE NOT NULL,
#             password TEXT NOT NULL
#         )
#     ''')
#     conn.commit()
#     conn.close()

# # Call the function to ensure the table exists
# init_users_db()
# @app.route('/')
# def index():
#     year_filter = request.args.get('year')
#     conn = get_db_connection()
    
#     # 1. Dynamically find all unique years in the database
#     # This regex handles formats like "2026", "Feb 2026", etc.
#     all_dates = conn.execute('SELECT DISTINCT date FROM events').fetchall()
#     years = set()
#     for row in all_dates:
#         # Extract 4-digit years using regex
#         found = re.findall(r'20\d{2}', str(row['date']))
#         for y in found:
#             years.add(y)
#     sorted_years = sorted(list(years)) # [2023, 2024, 2025, 2026]

#     # 2. Filter records based on selection
#     if year_filter and year_filter != "All":
#         query = '''
#             SELECT s.name, e.event_name, e.date, e.venue, p.remarks
#             FROM participation p
#             JOIN students s ON p.student_id = s.id
#             JOIN events e ON p.event_id = e.id
#             WHERE e.date LIKE ?
#             ORDER BY e.id DESC
#         '''
#         records = conn.execute(query, ('%' + year_filter + '%',)).fetchall()
#     else:
#         query = '''
#             SELECT s.name, e.event_name, e.date, e.venue, p.remarks
#             FROM participation p
#             JOIN students s ON p.student_id = s.id
#             JOIN events e ON p.event_id = e.id
#             ORDER BY e.id DESC
#         '''
#         records = conn.execute(query).fetchall()
        
#     conn.close()
#     return render_template('index.html', 
#                            records=records, 
#                            current_year=year_filter, 
#                            available_years=sorted_years)
# # @app.route('/')
# # def index():
# #     year = request.args.get('year') # Get the year from the URL (e.g., /?year=2026)
# #     conn = get_db_connection()
    
# #     if year and year != "All":
# #         # Filter records by the selected year
# #         query = '''
# #             SELECT s.name, e.event_name, e.date, e.venue, p.remarks
# #             FROM participation p
# #             JOIN students s ON p.student_id = s.id
# #             JOIN events e ON p.event_id = e.id
# #             WHERE e.date LIKE ?
# #             ORDER BY e.id DESC
# #         '''
# #         records = conn.execute(query, ('%' + year + '%',)).fetchall()
# #     else:
# #         # Default: Show everything
# #         query = '''
# #             SELECT s.name, e.event_name, e.date, e.venue, p.remarks
# #             FROM participation p
# #             JOIN students s ON p.student_id = s.id
# #             JOIN events e ON p.event_id = e.id
# #             ORDER BY e.id DESC
# #         '''
# #         records = conn.execute(query).fetchall()
        
# #     conn.close()
# #     return render_template('index.html', records=records, current_year=year)
# # @app.route('/')
# # def index():
# #     conn = get_db_connection()
# #     # Joined query to get Student Name, Event details, and Remarks
# #     query = '''
# #         SELECT s.name, e.event_name, e.date, e.venue, p.remarks
# #         FROM participation p
# #         JOIN students s ON p.student_id = s.id
# #         JOIN events e ON p.event_id = e.id
# #         ORDER BY e.id DESC
# #     '''
# #     records = conn.execute(query).fetchall()
# #     conn.close()
# #     return render_template('index.html', records=records)

# @app.route('/leaderboard')
# def leaderboard():
#     conn = get_db_connection()
#     # Query to count participations per student
#     query = '''
#         SELECT s.name, COUNT(p.id) as count
#         FROM students s
#         JOIN participation p ON s.id = p.student_id
#         GROUP BY s.name
#         ORDER BY count DESC
#         LIMIT 10
#     '''
#     top_students = conn.execute(query).fetchall()
#     conn.close()
#     return render_template('leaderboard.html', top_students=top_students)

# if __name__ == '__main__':
#     # Run the server
#     print("Starting server at http://127.0.0.1:5000")
#     app.run(debug=True)

import os

from flask import Flask, redirect, render_template, request, session, url_for, flash, send_file
import sqlite3
import re
import io
from docx import Document

app = Flask(__name__)
app.secret_key = 'secret_key_for_session_and_flashing' # Required for flash()

# This line finds the folder where your app.py is sitting on Render
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
DB_PATH = os.path.join(BASE_DIR, "hackathon.db")

# --- DATABASE HELPERS ---

def get_db_connection():
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    return conn

def init_users_db():
    conn = get_db_connection()
    conn.execute('''
        CREATE TABLE IF NOT EXISTS users (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            fullname TEXT NOT NULL,
            email TEXT UNIQUE NOT NULL,
            password TEXT NOT NULL
        )
    ''')
    conn.commit()
    conn.close()

# Initialize the user table on startup
init_users_db()

with app.app_context():
    db = get_db_connection()
    # Create the table if it doesn't exist
    db.execute('CREATE TABLE IF NOT EXISTS users (id INTEGER PRIMARY KEY, fullname TEXT, email TEXT UNIQUE, password TEXT)')
    
    # Add a MASTER ACCOUNT for your demo
    try:
        db.execute('INSERT INTO users (fullname, email, password) VALUES (?, ?, ?)', 
                   ('Admin User', 'admin@test.com', 'admin123'))
        db.commit()
    except:
        pass # Email already exists, no need to add again
    db.close()

# --- AUTHENTICATION ROUTES ---

@app.route('/')
def home():
    """Redirects the root URL to the login page."""
    return render_template('login.html')

@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        email = request.form.get('email')
        password = request.form.get('password')

        conn = get_db_connection()
        user = conn.execute('SELECT * FROM users WHERE email = ?', (email,)).fetchone()
        conn.close()

        if user and user['password'] == password:
            # THIS IS THE PART YOU ARE MISSING:
            session['user_id'] = user['id']
            return redirect(url_for('dashboard'))
        else:
            flash("Invalid email or password.")
            return render_template('login.html')
            
    return render_template('login.html')

@app.route('/signup', methods=['GET', 'POST'])
def signup():
    if request.method == 'POST':
        fullname = request.form.get('fullname')
        email = request.form.get('email')
        password = request.form.get('password')
        confirm_password = request.form.get('confirm_password')

        if password != confirm_password:
            flash("Passwords do not match!")
            return redirect(url_for('signup'))

        if len(password) < 8:
            flash("Password must be at least 8 characters.")
            return redirect(url_for('signup'))

        conn = get_db_connection()
        try:
            conn.execute('INSERT INTO users (fullname, email, password) VALUES (?, ?, ?)',
                         (fullname, email, password))
            conn.commit()
            conn.close() # Close it here
            flash("Account created! Please login.")
            return redirect(url_for('login')) # Redirect immediately after success
            # flash("Account created! Please login.")
            # return redirect(url_for('login'))
        except sqlite3.IntegrityError:
            # flash("Email already exists!")
            flash("Email already exists!")
            return redirect(url_for('signup'))
        except Exception as e:
            print(f"Error: {e}") # This will show in Render logs if it fails
            return "Database Error", 500
        # finally:
        #     conn.close()
            
    return render_template('signup.html')

# --- MAIN DASHBOARD & CRUD ---

@app.route('/dashboard')
def dashboard():
    year_filter = request.args.get('year')
    conn = get_db_connection()
    if 'user_id' not in session:
        return redirect(url_for('login'))
    # NEW: Fetch the logged-in user's name
    user = conn.execute('SELECT fullname FROM users WHERE id = ?', (session['user_id'],)).fetchone()
    user_name = user['fullname'] if user else "Admin"

    # 1. Get unique years for the filter buttons
    all_dates = conn.execute('SELECT DISTINCT date FROM events').fetchall()
    years = set()
    for row in all_dates:
        found = re.findall(r'20\d{2}', str(row['date']))
        for y in found:
            years.add(y)
    sorted_years = sorted(list(years))

    # 2. Build the query based on filter
    if year_filter and year_filter != "All":
        # Updated Query (Notice we use p.remarks now)
        query = """
            SELECT p.id, s.name, e.event_name, e.date, e.venue, e.prizes, COALESCE(p.remarks, e.prizes) AS final_prize 
            FROM participation p
            JOIN students s ON p.student_id = s.id
            JOIN events e ON p.event_id = e.id
            WHERE e.date LIKE ?
            ORDER BY p.id DESC
        """
        # query = '''
        #     SELECT p.id, s.name, e.event_name, e.date, e.venue, p.prizes,COALESCE(p.remarks, e.prizes) AS final_prize
        #     FROM participation p
        #     JOIN students s ON p.student_id = s.id
        #     JOIN events e ON p.event_id = e.id
        #     WHERE e.date LIKE ?
        #     ORDER BY p.id DESC
        # '''
        records = conn.execute(query, ('%' + year_filter + '%',)).fetchall()
    else:
        # Updated Query (Notice we use p.remarks now)
        # query = """
        #     SELECT p.id, s.name, e.event_name, e.date, e.venue, e.prizes, COALESCE(p.remarks, e.prizes) AS final_prize 
        #     FROM participation p
        #     JOIN students s ON p.student_id = s.id
        #     JOIN events e ON p.event_id = e.id
        #     WHERE e.date LIKE ?
        #     ORDER BY p.id DESC
        # """
        # No 'WHERE' clause here, so no '?' is needed
        query = """
            SELECT p.id, s.name, e.event_name, e.date, e.venue, e.prizes, COALESCE(p.remarks, e.prizes) AS final_prize 
            FROM participation p
            JOIN students s ON p.student_id = s.id
            JOIN events e ON p.event_id = e.id
            ORDER BY p.id DESC
        """
        records = conn.execute(query).fetchall()
        
    conn.close()
    return render_template('index.html', 
                           records=records, 
                           current_year=year_filter, 
                           available_years=sorted_years,
                           user_name=user_name)

@app.route('/delete-record/<int:id>')
def delete_record(id):
    conn = get_db_connection()
    conn.execute('DELETE FROM participation WHERE id = ?', (id,))
    conn.commit()
    conn.close()
    flash("Record deleted successfully!")
    return redirect(url_for('dashboard'))

@app.route('/export-word')
def export_word():
    conn = get_db_connection()
    records = conn.execute('''
        SELECT s.name, e.event_name, e.date, e.venue 
        FROM participation p 
        JOIN students s ON p.student_id = s.id 
        JOIN events e ON p.event_id = e.id
    ''').fetchall()
    conn.close()

    doc = Document()
    doc.add_heading('BITM CSE - Hackathon Participation Report', 0)

    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Student Name'
    hdr_cells[1].text = 'Event Name'
    hdr_cells[2].text = 'Date'
    hdr_cells[3].text = 'Venue'

    for rec in records:
        row_cells = table.add_row().cells
        row_cells[0].text = str(rec['name'])
        row_cells[1].text = str(rec['event_name'])
        row_cells[2].text = str(rec['date'])
        row_cells[3].text = str(rec['venue'])

    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return send_file(file_stream, as_attachment=True, download_name='Hackathon_Report.docx')

@app.route('/leaderboard')
def leaderboard():
    conn = get_db_connection()
    query = '''
        SELECT s.name, COUNT(p.id) as count
        FROM students s
        JOIN participation p ON s.id = p.student_id
        GROUP BY s.name
        ORDER BY count DESC
        LIMIT 10
    '''
    top_students = conn.execute(query).fetchall()
    conn.close()
    return render_template('leaderboard.html', top_students=top_students)

@app.route('/logout')
def logout():
    # In a real app, you'd use session.clear() here
    flash("You have been logged out.")
    return redirect(url_for('home'))

@app.route('/add-record', methods=['POST'])
def add_record():
    student_name = request.form.get('student_name')
    event_name = request.form.get('event_name')
    date = request.form.get('date')
    venue = request.form.get('venue')
    prizes = request.form.get('prizes') # New line to get the data
    conn = get_db_connection()
    
    # 1. Handle Student: Check if they exist, or create new
    conn.execute('INSERT OR IGNORE INTO students (name) VALUES (?)', (student_name,))
    student = conn.execute('SELECT id FROM students WHERE name = ?', (student_name,)).fetchone()
    student_id = student['id']
    
    # 2. Handle Event: Create the event entry
    # conn.execute('INSERT INTO events (event_name, date, venue) VALUES (?, ?, ?)', 
    #              (event_name, date, venue))
    # Update your INSERT statement to include 'prizes'
    conn.execute('INSERT INTO events (event_name, date, venue, prizes) VALUES (?, ?, ?, ?)', 
             (event_name, date, venue, prizes))
    event_id = conn.execute('SELECT last_insert_rowid()').fetchone()[0]
    
    # 3. Link them in Participation
    # conn.execute('INSERT INTO participation (student_id, event_id) VALUES (?, ?)', 
    #              (student_id, event_id))
    
    # 2. Save to Participation table (so they match!)
    conn.execute('INSERT INTO participation (student_id, event_id, remarks) VALUES (?, ?, ?)', 
                 (student_id, event_id, prizes)) # Save the prizes in the remarks column for now
    
    conn.commit()
    conn.close()
    
    flash("New record added successfully!")
    return redirect(url_for('dashboard'))

# if __name__ == '__main__':
#     app.run(debug=True)
if __name__ == "__main__":
    # Render provides a 'PORT' variable, otherwise use 5000
    port = int(os.environ.get("PORT", 5000))
    app.run(host='0.0.0.0', port=port)

# # This ensures the 'users' table exists as soon as the app starts
# with app.app_context():
#     db = get_db_connection()
#     db.execute('''
#         CREATE TABLE IF NOT EXISTS users (
#             id INTEGER PRIMARY KEY AUTOINCREMENT,
#             fullname TEXT NOT NULL,
#             email TEXT UNIQUE NOT NULL,
#             password TEXT NOT NULL
#         )
#     ''')
#     db.commit()
#     db.close()
# Create a permanent account for the demo
with app.app_context():
    conn = get_db_connection()
    # Ensure table exists
    conn.execute('''CREATE TABLE IF NOT EXISTS users 
                 (id INTEGER PRIMARY KEY AUTOINCREMENT, fullname TEXT, email TEXT UNIQUE, password TEXT)''')
    # Add a demo user (ignores if already exists)
    try:
        conn.execute('INSERT INTO users (fullname, email, password) VALUES (?, ?, ?)',
                     ('Admin', 'admin@test.com', 'admin123'))
        conn.commit()
    except:
        pass
    conn.close()