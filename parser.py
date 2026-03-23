import sqlite3
import os
import re
from docx import Document

def clean_data(event_text):
    """
    Specifically targets 'Cash Prize' and currency patterns to move them 
    to the rewards column and remove them from the event title.
    """
    # 1. First, check for brackets and extract content inside them
    bracket_pattern = r"\((.*?)\)"
    bracket_match = re.search(bracket_pattern, event_text)
    
    bracket_remarks = ""
    if bracket_match:
        bracket_remarks = bracket_match.group(1).strip()
        event_text = re.sub(bracket_pattern, "", event_text).strip()

    # 2. UPDATED: Improved pattern to catch "Cash Prize", "Rs.", and "/-"
    # This will now catch "Cash Prize Rs. 20,000/-" perfectly.
    prize_pattern = r"(Cash Prize.*|Rs.\..*|Winner.*|.*Prize.*|\d+.*Place.*|.*Cash.*|.*/-.*)"
    keyword_match = re.search(prize_pattern, event_text, re.IGNORECASE)
    
    keyword_remarks = ""
    if keyword_match:
        keyword_remarks = keyword_match.group(0).strip()
        # Remove the prize info from the event name
        event_text = event_text.replace(keyword_remarks, "").strip().rstrip(',')

    # Combine and remove any leftover brackets
    final_remarks = f"{bracket_remarks} {keyword_remarks}".strip().replace('(', '').replace(')', '')
    return event_text.strip(), final_remarks

def import_all_docs():
    data_folder = 'data'
    # Start with a fresh database
    if os.path.exists('hackathon.db'):
        os.remove('hackathon.db')
    
    conn = sqlite3.connect('hackathon.db')
    cursor = conn.cursor()
    cursor.execute('CREATE TABLE students (id INTEGER PRIMARY KEY AUTOINCREMENT, name TEXT UNIQUE)')
    cursor.execute('CREATE TABLE events (id INTEGER PRIMARY KEY AUTOINCREMENT, event_name TEXT, date TEXT, venue TEXT)')
    cursor.execute('CREATE TABLE participation (id INTEGER PRIMARY KEY AUTOINCREMENT, student_id INTEGER, event_id INTEGER, remarks TEXT)')

    files = [f for f in os.listdir(data_folder) if f.endswith('.docx')]

    for file_name in files:
        file_path = os.path.join(data_folder, file_name)
        doc = Document(file_path)
        print(f"Cleaning cash prizes and importing {file_name}...")

        for table in doc.tables:
            for row in table.rows[1:]:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if len(cells) < 3: continue

                name = cells[1]
                raw_event = cells[2]
                
                # Cleaning step that now handles Cash Prize relocation
                event_name, remarks = clean_data(raw_event)

                # Column Shift Handling
                date = ""
                venue = ""
                remaining_cells = cells[3:]
                if len(remaining_cells) >= 2:
                    date = remaining_cells[0]
                    venue = remaining_cells[1]
                elif len(remaining_cells) == 1:
                    date = remaining_cells[0]

                # Special case check for 2026 docs 
                if len(cells) > 5 and not remarks:
                    remarks = cells[4].replace('(', '').replace(')', '').strip()
                elif remarks:
                    remarks = remarks.replace('(', '').replace(')', '').strip()

                # Database Insertion
                cursor.execute("INSERT OR IGNORE INTO students (name) VALUES (?)", (name,))
                cursor.execute("SELECT id FROM students WHERE name = ?", (name,))
                s_id = cursor.fetchone()[0]

                cursor.execute("INSERT INTO events (event_name, date, venue) VALUES (?, ?, ?)", 
                               (event_name, date, venue))
                e_id = cursor.lastrowid

                cursor.execute("INSERT INTO participation (student_id, event_id, remarks) VALUES (?, ?, ?)",
                               (s_id, e_id, remarks))

    conn.commit()
    conn.close()
    print("✔ SUCCESS: All Cash Prizes relocated and brackets removed!")

if __name__ == "__main__":
    import_all_docs()
# import sqlite3
# import os
# import re
# from docx import Document

# def clean_data(event_text):
#     """
#     Extracts prize/winner info from event text, including text inside brackets.
#     Removes the brackets () from the final output.
#     """
#     # 1. First, check for brackets and extract content inside them
#     bracket_pattern = r"\((.*?)\)"
#     bracket_match = re.search(bracket_pattern, event_text)
    
#     bracket_remarks = ""
#     if bracket_match:
#         bracket_remarks = bracket_match.group(1).strip() # Get content inside ()
#         event_text = re.sub(bracket_pattern, "", event_text).strip() # Remove the () block

#     # 2. Check for other prize keywords (Winner, Prize, Rs., etc.)
#     prize_pattern = r"(Winner.*|.*Prize.*|Rs\..*|\d+.*Place.*)"
#     keyword_match = re.search(prize_pattern, event_text, re.IGNORECASE)
    
#     keyword_remarks = ""
#     if keyword_match:
#         keyword_remarks = keyword_match.group(0).strip()
#         event_text = event_text.replace(keyword_remarks, "").strip().rstrip(',')

#     # Combine both types of remarks found
#     final_remarks = f"{bracket_remarks} {keyword_remarks}".strip()
#     return event_text, final_remarks

# def import_all_docs():
#     data_folder = 'data'
#     if os.path.exists('hackathon.db'):
#         os.remove('hackathon.db')
    
#     conn = sqlite3.connect('hackathon.db')
#     cursor = conn.cursor()
#     cursor.execute('CREATE TABLE students (id INTEGER PRIMARY KEY AUTOINCREMENT, name TEXT UNIQUE)')
#     cursor.execute('CREATE TABLE events (id INTEGER PRIMARY KEY AUTOINCREMENT, event_name TEXT, date TEXT, venue TEXT)')
#     cursor.execute('CREATE TABLE participation (id INTEGER PRIMARY KEY AUTOINCREMENT, student_id INTEGER, event_id INTEGER, remarks TEXT)')

#     files = [f for f in os.listdir(data_folder) if f.endswith('.docx')]

#     for file_name in files:
#         file_path = os.path.join(data_folder, file_name)
#         doc = Document(file_path)
#         print(f"Cleaning brackets and importing {file_name}...")

#         for table in doc.tables:
#             for row in table.rows[1:]:
#                 cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
#                 if len(cells) < 3: continue

#                 name = cells[1]
#                 raw_event = cells[2]
                
#                 # CLEANING STEP: This now removes brackets from both columns
#                 event_name, remarks = clean_data(raw_event)

#                 # Column Shift Handling
#                 date = ""
#                 venue = ""
#                 remaining_cells = cells[3:]
#                 if len(remaining_cells) >= 2:
#                     date = remaining_cells[0]
#                     venue = remaining_cells[1]
#                 elif len(remaining_cells) == 1:
#                     date = remaining_cells[0]

#                 # Special case for 2026 docs where remarks might be in cell 5
#                 if len(cells) > 5 and not remarks:
#                     # Strip any brackets found in cell 5 as well
#                     remarks = cells[4].replace('(', '').replace(')', '').strip()
#                 elif remarks:
#                     # Final safety strip of any leftover brackets
#                     remarks = remarks.replace('(', '').replace(')', '').strip()

#                 # Database Insertion
#                 cursor.execute("INSERT OR IGNORE INTO students (name) VALUES (?)", (name,))
#                 cursor.execute("SELECT id FROM students WHERE name = ?", (name,))
#                 s_id = cursor.fetchone()[0]

#                 cursor.execute("INSERT INTO events (event_name, date, venue) VALUES (?, ?, ?)", 
#                                (event_name, date, venue))
#                 e_id = cursor.lastrowid

#                 cursor.execute("INSERT INTO participation (student_id, event_id, remarks) VALUES (?, ?, ?)",
#                                (s_id, e_id, remarks))

#     conn.commit()
#     conn.close()
#     print("✔ Data cleaned (brackets removed) and re-imported successfully!")

# if __name__ == "__main__":
#     import_all_docs()

# import sqlite3
# import os
# import re
# from docx import Document

# def clean_data(event_text):
#     """
#     Specifically finds 'Cash Prize' and currency info to move to Remarks.
#     Removes brackets () and ensures column alignment.
#     """
#     # 1. Look for Cash Prize info specifically (e.g., Cash Prize Rs. 20,000/-)
#     # This regex looks for 'Cash Prize' followed by anything until the end of a line or punctuation
#     cash_prize_pattern = r"(Cash Prize\s*Rs\..*|Winner.*|.*Prize.*|\d+\s*(?:st|nd|rd|th)\s*Place.*)"
#     cash_match = re.search(cash_prize_pattern, event_text, re.IGNORECASE)
    
#     found_remarks = ""
#     if cash_match:
#         found_remarks = cash_match.group(0).strip()
#         # Remove the prize string from the event name
#         event_text = event_text.replace(found_remarks, "").strip()

#     # 2. Look for text inside brackets
#     bracket_pattern = r"\((.*?)\)"
#     bracket_match = re.search(bracket_pattern, event_text)
    
#     if bracket_match:
#         bracket_text = bracket_match.group(1).strip()
#         found_remarks = f"{bracket_text} {found_remarks}".strip()
#         event_text = re.sub(bracket_pattern, "", event_text).strip()

#     # Clean up any trailing commas or dashes from the event name
#     clean_event = event_text.rstrip(',').rstrip('-').strip()
#     clean_remarks = found_remarks.replace('(', '').replace(')', '').strip()
    
#     return clean_event, clean_remarks

# def import_all_docs():
#     data_folder = 'data'
#     # Start fresh to ensure no old messy data remains
#     if os.path.exists('hackathon.db'):
#         os.remove('hackathon.db')
    
#     conn = sqlite3.connect('hackathon.db')
#     cursor = conn.cursor()
#     cursor.execute('CREATE TABLE students (id INTEGER PRIMARY KEY AUTOINCREMENT, name TEXT UNIQUE)')
#     cursor.execute('CREATE TABLE events (id INTEGER PRIMARY KEY AUTOINCREMENT, event_name TEXT, date TEXT, venue TEXT)')
#     cursor.execute('CREATE TABLE participation (id INTEGER PRIMARY KEY AUTOINCREMENT, student_id INTEGER, event_id INTEGER, remarks TEXT)')

#     files = [f for f in os.listdir(data_folder) if f.endswith('.docx')]

#     for file_name in files:
#         file_path = os.path.join(data_folder, file_name)
#         doc = Document(file_path)
#         print(f"🔄 Relocating Cash Prizes in: {file_name}")

#         for table in doc.tables:
#             for row in table.rows[1:]:
#                 cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
#                 if len(cells) < 3: continue

#                 name = cells[1]
#                 raw_event = cells[2]
                
#                 # EXECUTE FIX: This separates the Cash Prize from the name
#                 event_name, remarks = clean_data(raw_event)

#                 # Fix Column Shifts
#                 remaining = cells[3:]
#                 date = ""
#                 venue = ""
                
#                 if len(remaining) >= 2:
#                     date, venue = remaining[0], remaining[1]
#                 elif len(remaining) == 1:
#                     date = remaining[0]

#                 # Special check for 2026 docs where prizes are sometimes in cell 5
#                 if len(cells) > 5 and not remarks:
#                     remarks = cells[4].replace('(', '').replace(')', '').strip()

#                 # Database Insertion
#                 cursor.execute("INSERT OR IGNORE INTO students (name) VALUES (?)", (name,))
#                 cursor.execute("SELECT id FROM students WHERE name = ?", (name,))
#                 s_id = cursor.fetchone()[0]

#                 cursor.execute("INSERT INTO events (event_name, date, venue) VALUES (?, ?, ?)", 
#                                (event_name, date, venue))
#                 e_id = cursor.lastrowid

#                 cursor.execute("INSERT INTO participation (student_id, event_id, remarks) VALUES (?, ?, ?)",
#                                (s_id, e_id, remarks))

#     conn.commit()
#     conn.close()
#     print("✔ SUCCESS: All Cash Prizes moved to the Remarks/Prizes column!")

# if __name__ == "__main__":
#     import_all_docs()